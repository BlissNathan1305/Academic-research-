import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns

# -------------------------------
# Function: set single line spacing
# -------------------------------
def set_single_spacing(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(ns.qn('w:line'), "240")  # single spacing
    spacing.set(ns.qn('w:lineRule'), "auto")
    pPr.append(spacing)

# -------------------------------
# Process one document
# -------------------------------
def process_docx(file_path):
    doc = Document(file_path)

    # Page setup
    section = doc.sections[0]
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Header alignment (top right)
    header = section.header
    if header.paragraphs:
        header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Format tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    set_single_spacing(para)

                    for run in para.runs:
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(0, 0, 0)

    # Ensure all body text is black
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)

    # Save formatted copy
    base, ext = os.path.splitext(file_path)
    output_file = f"{base}_formatted{ext}"
    doc.save(output_file)

    print(f"✔ Formatted: {output_file}")

# -------------------------------
# Scan current folder
# -------------------------------
current_folder = os.getcwd()

for filename in os.listdir(current_folder):
    if (
        filename.lower().endswith(".docx")
        and not filename.startswith("~$")
        and not filename.endswith("_formatted.docx")
    ):
        process_docx(os.path.join(current_folder, filename))
